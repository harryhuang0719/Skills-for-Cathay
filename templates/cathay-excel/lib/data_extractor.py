"""
Cathay Capital PE Financial Model — Data Extractor.

Extracts financial data from source materials (xlsx, docx, md) in a folder
and produces a standardized data_dict compatible with model_populator.

Usage:
    from data_extractor import extract_from_folder
    data = extract_from_folder("~/Desktop/340-蔚蓝支点/")
"""

import sys
import os
import re
import glob
from pathlib import Path

_LIB_DIR = os.path.dirname(os.path.abspath(__file__))
if _LIB_DIR not in sys.path:
    sys.path.insert(0, _LIB_DIR)


# =============================================================================
# Constants
# =============================================================================

# Chinese financial text patterns
PATTERNS = {
    'revenue': r'收入[：:]\s*([\d,.]+)\s*(万|百万|千万|亿)',
    'gross_margin': r'毛利率[：:]\s*([\d.]+)%',
    'net_income': r'净利润?[：:]\s*([\d,.]+)\s*(万|百万|千万|亿)',
    'headcount': r'员工[：:]\s*(\d+)\s*人',
    'capex': r'[Cc]ap[Ee]x[：:]\s*([\d,.]+)',
    'valuation': r'估值[：:]\s*([\d,.]+)\s*(万|百万|千万|亿)',
    'revenue_growth': r'收入增[长速][：:]\s*([\d.]+)%',
    'ebitda': r'EBITDA[：:]\s*([\d,.]+)\s*(万|百万|千万|亿)',
    'debt': r'负债[：:]\s*([\d,.]+)\s*(万|百万|千万|亿)',
    'cash': r'现金[：:]\s*([\d,.]+)\s*(万|百万|千万|亿)',
    'ar_days': r'应收[账帐]款周转[天日][数：:]\s*([\d.]+)',
    'inventory_days': r'存货周转[天日][数：:]\s*([\d.]+)',
    'ap_days': r'应付[账帐]款周转[天日][数：:]\s*([\d.]+)',
}

# Keywords that indicate financial statement sheets
PL_KEYWORDS = ['利润', 'income', 'p&l', 'profit', 'loss', '损益']
BS_KEYWORDS = ['资产', 'balance', 'sheet', '负债', 'liability']
CF_KEYWORDS = ['现金', 'cash', 'flow', '流量']
REVENUE_KEYWORDS = ['收入', 'revenue', 'sales', '营业']


# =============================================================================
# Public API
# =============================================================================

def extract_from_folder(folder_path):
    """Scan a folder for financial source materials and extract data.

    Looks for:
    - .xlsx files: searches for sheets with financial data (P&L, BS, CF keywords)
    - .docx files: extracts tables with financial data
    - .md files: extracts numbers with context using regex

    Args:
        folder_path: path to folder containing source materials

    Returns:
        data_dict compatible with model_populator.populate_model()
        with 'confidence' scores for extracted values
    """
    folder = Path(os.path.expanduser(folder_path))
    if not folder.is_dir():
        raise ValueError(f"Folder not found: {folder}")

    extractions = []

    # Scan for source files
    xlsx_files = list(folder.glob('**/*.xlsx'))
    docx_files = list(folder.glob('**/*.docx'))
    md_files = list(folder.glob('**/*.md'))

    for f in xlsx_files:
        try:
            ext = extract_from_xlsx(f)
            if ext:
                extractions.append(('xlsx', ext))
        except Exception as e:
            print(f"Warning: Failed to extract from {f.name}: {e}")

    for f in docx_files:
        try:
            ext = extract_from_docx(f)
            if ext:
                extractions.append(('docx', ext))
        except Exception as e:
            print(f"Warning: Failed to extract from {f.name}: {e}")

    for f in md_files:
        try:
            ext = extract_from_md(f)
            if ext:
                extractions.append(('md', ext))
        except Exception as e:
            print(f"Warning: Failed to extract from {f.name}: {e}")

    return merge_extractions(extractions)


# =============================================================================
# Excel Extraction
# =============================================================================

def extract_from_xlsx(path):
    """Extract financial data from an Excel file.

    Scans sheet names for financial keywords, then extracts tabular data
    from matching sheets.

    Args:
        path: Path to .xlsx file

    Returns:
        dict with extracted financial data, or None if no financial data found
    """
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    result = {'historical': {}, 'source': str(path)}

    for sheet_name in wb.sheetnames:
        name_lower = sheet_name.lower()

        if _matches_keywords(name_lower, PL_KEYWORDS):
            _extract_pl_sheet(wb[sheet_name], result)
        elif _matches_keywords(name_lower, BS_KEYWORDS):
            _extract_bs_sheet(wb[sheet_name], result)
        elif _matches_keywords(name_lower, CF_KEYWORDS):
            _extract_cf_sheet(wb[sheet_name], result)
        elif _matches_keywords(name_lower, REVENUE_KEYWORDS):
            _extract_revenue_sheet(wb[sheet_name], result)

    wb.close()
    return result if result['historical'] else None


def _matches_keywords(text, keywords):
    """Check if text contains any keyword."""
    return any(kw in text for kw in keywords)


def _extract_pl_sheet(ws, result):
    """Extract P&L data from a worksheet."""
    years, row_data = _parse_financial_table(ws)
    if not years:
        return

    # Map row labels to data_dict keys
    pl_label_map = {
        '收入': 'revenue', '营业收入': 'revenue', 'revenue': 'revenue',
        '营业成本': 'cogs', '成本': 'cogs', 'cogs': 'cogs',
        'cost of': 'cogs',
        '销售费用': 'sga', '管理费用': 'sga', 'sg&a': 'sga',
        '折旧': 'da', '摊销': 'da', 'depreciation': 'da', 'd&a': 'da',
        '利息': 'interest', 'interest': 'interest',
        '所得税': 'tax', '税': 'tax', 'tax': 'tax', 'income tax': 'tax',
        '净利润': 'net_income', 'net income': 'net_income', 'net profit': 'net_income',
    }

    sga_components = []

    for label, values in row_data:
        label_lower = label.lower().strip()
        matched_key = None

        for pattern, key in pl_label_map.items():
            if pattern in label_lower:
                matched_key = key
                break

        if matched_key:
            for i, year in enumerate(years):
                if i < len(values) and values[i] is not None:
                    if year not in result['historical']:
                        result['historical'][year] = {}
                    # SG&A may appear as separate lines; accumulate
                    if matched_key == 'sga' and '费用' in label:
                        sga_components.append((year, i, values[i]))
                    else:
                        result['historical'][year][matched_key] = _to_number(values[i])

    # Sum SGA components
    if sga_components:
        sga_by_year = {}
        for year, i, val in sga_components:
            sga_by_year[year] = sga_by_year.get(year, 0) + _to_number(val)
        for year, total in sga_by_year.items():
            if year not in result['historical']:
                result['historical'][year] = {}
            result['historical'][year]['sga'] = total


def _extract_bs_sheet(ws, result):
    """Extract Balance Sheet data from a worksheet."""
    years, row_data = _parse_financial_table(ws)
    if not years:
        return

    bs_label_map = {
        '现金': 'cash', '货币资金': 'cash', 'cash': 'cash',
        '应收': 'ar', 'receivable': 'ar', 'accounts receivable': 'ar',
        '存货': 'inventory', 'inventory': 'inventory',
        '固定资产': 'ppe', 'ppe': 'ppe', 'property': 'ppe',
        '应付': 'ap', 'payable': 'ap', 'accounts payable': 'ap',
        '借款': 'debt', '负债': 'debt', 'debt': 'debt', 'loan': 'debt',
        '股本': 'equity', '实收资本': 'equity', 'equity': 'equity',
        '留存': 'retained_earnings', '未分配利润': 'retained_earnings',
        'retained': 'retained_earnings',
    }

    for label, values in row_data:
        label_lower = label.lower().strip()
        matched_key = None

        for pattern, key in bs_label_map.items():
            if pattern in label_lower:
                matched_key = key
                break

        if matched_key:
            for i, year in enumerate(years):
                if i < len(values) and values[i] is not None:
                    if year not in result['historical']:
                        result['historical'][year] = {}
                    result['historical'][year][matched_key] = _to_number(values[i])


def _extract_cf_sheet(ws, result):
    """Extract Cash Flow data from a worksheet."""
    years, row_data = _parse_financial_table(ws)
    if not years:
        return

    cf_label_map = {
        '资本支出': 'capex', 'capex': 'capex', '购建': 'capex',
        '购置': 'capex', 'capital expenditure': 'capex',
    }

    for label, values in row_data:
        label_lower = label.lower().strip()
        matched_key = None

        for pattern, key in cf_label_map.items():
            if pattern in label_lower:
                matched_key = key
                break

        if matched_key:
            for i, year in enumerate(years):
                if i < len(values) and values[i] is not None:
                    if year not in result['historical']:
                        result['historical'][year] = {}
                    result['historical'][year][matched_key] = abs(_to_number(values[i]))


def _extract_revenue_sheet(ws, result):
    """Extract revenue breakdown data from a worksheet."""
    years, row_data = _parse_financial_table(ws)
    if not years:
        return

    for label, values in row_data:
        label_lower = label.lower().strip()
        if any(kw in label_lower for kw in ['总计', 'total', '合计', '收入']):
            for i, year in enumerate(years):
                if i < len(values) and values[i] is not None:
                    if year not in result['historical']:
                        result['historical'][year] = {}
                    if 'revenue' not in result['historical'][year]:
                        result['historical'][year]['revenue'] = _to_number(values[i])


def _parse_financial_table(ws):
    """Parse a financial table worksheet.

    Looks for a year header row, then extracts label + value rows below.

    Returns:
        (years: list[int], row_data: list[(label, values)])
    """
    years = []
    row_data = []
    year_row_idx = None

    # Find the year header row (contains 4-digit numbers 2018-2030)
    for r_idx in range(1, min(ws.max_row + 1, 20)):
        row_values = []
        for c_idx in range(1, min(ws.max_column + 1, 20)):
            val = ws.cell(row=r_idx, column=c_idx).value
            row_values.append(val)

        # Check if this row contains year numbers
        found_years = []
        year_cols = []
        for c_idx, val in enumerate(row_values):
            year_int = _try_parse_year(val)
            if year_int and 2015 <= year_int <= 2035:
                found_years.append(year_int)
                year_cols.append(c_idx)

        if len(found_years) >= 2:
            years = found_years
            year_row_idx = r_idx
            break

    if not year_row_idx:
        return [], []

    # Extract data rows below the year header
    for r_idx in range(year_row_idx + 1, ws.max_row + 1):
        label_cell = ws.cell(row=r_idx, column=1).value
        if not label_cell or not isinstance(label_cell, str):
            # Try column 2 if column 1 is empty
            label_cell = ws.cell(row=r_idx, column=2).value
            if not label_cell or not isinstance(label_cell, str):
                continue

        label = str(label_cell).strip()
        if not label:
            continue

        values = []
        for c_idx in year_cols:
            val = ws.cell(row=r_idx, column=c_idx + 1).value  # c_idx is 0-based
            values.append(val)

        row_data.append((label, values))

    return years, row_data


# =============================================================================
# Word Document Extraction
# =============================================================================

def extract_from_docx(path):
    """Extract financial data from a Word document.

    Searches all tables for financial data patterns.

    Args:
        path: Path to .docx file

    Returns:
        dict with extracted financial data, or None if no financial data found
    """
    from docx import Document

    doc = Document(path)
    result = {'historical': {}, 'source': str(path)}

    # Extract data from tables
    for table in doc.tables:
        _extract_from_table(table, result)

    # Extract data from paragraphs (regex patterns)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    text_data = _extract_from_text(full_text)
    _merge_text_data(result, text_data)

    return result if result['historical'] or text_data else None


def _extract_from_table(table, result):
    """Extract financial data from a docx table."""
    rows = table.rows
    if len(rows) < 2:
        return

    # Check first row for years
    header_cells = [cell.text.strip() for cell in rows[0].cells]
    years = []
    year_indices = []

    for i, cell_text in enumerate(header_cells):
        year_int = _try_parse_year(cell_text)
        if year_int and 2015 <= year_int <= 2035:
            years.append(year_int)
            year_indices.append(i)

    if len(years) < 2:
        return

    # Process data rows
    label_map = {
        '收入': 'revenue', '营业收入': 'revenue', 'revenue': 'revenue',
        '营业成本': 'cogs', '成本': 'cogs',
        '毛利': 'gross_profit',
        '净利润': 'net_income', 'net income': 'net_income',
        '现金': 'cash', '货币资金': 'cash',
        '应收': 'ar',
        '存货': 'inventory',
        '固定资产': 'ppe',
        '应付': 'ap',
        '负债': 'debt',
        '资本支出': 'capex',
    }

    for row in rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if not cells:
            continue

        label = cells[0].lower()
        matched_key = None
        for pattern, key in label_map.items():
            if pattern in label:
                matched_key = key
                break

        if matched_key:
            for idx, year in zip(year_indices, years):
                if idx < len(cells):
                    val = _parse_number_text(cells[idx])
                    if val is not None:
                        if year not in result['historical']:
                            result['historical'][year] = {}
                        result['historical'][year][matched_key] = val


# =============================================================================
# Markdown Extraction
# =============================================================================

def extract_from_md(path):
    """Extract financial data from a Markdown file.

    Uses regex patterns to find Chinese financial text patterns.

    Args:
        path: Path to .md file

    Returns:
        dict with extracted financial data, or None if no financial data found
    """
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()

    text_data = _extract_from_text(text)

    if not text_data:
        return None

    result = {'historical': {}, 'text_data': text_data, 'source': str(path)}

    # Try to associate extracted values with years from context
    # Look for year mentions near financial data
    year_pattern = r'(20[12]\d)[年E]?'
    year_mentions = re.findall(year_pattern, text)

    # If we find a dominant year, assign values to it
    if year_mentions:
        # Use the most recent year mentioned
        latest_year = max(int(y) for y in year_mentions)
        if latest_year not in result['historical']:
            result['historical'][latest_year] = {}

        if 'revenue' in text_data:
            result['historical'][latest_year]['revenue'] = text_data['revenue']
        if 'net_income' in text_data:
            result['historical'][latest_year]['net_income'] = text_data['net_income']
        if 'cash' in text_data:
            result['historical'][latest_year]['cash'] = text_data['cash']

    return result if (result['historical'] or text_data) else None


def _extract_from_text(text):
    """Extract financial values from text using regex patterns.

    Returns:
        dict of extracted values (key -> value in RMB millions)
    """
    extracted = {}

    for key, pattern in PATTERNS.items():
        matches = re.findall(pattern, text)
        if matches:
            # Take the last match (usually the most relevant/recent)
            match = matches[-1]
            if key == 'gross_margin' or key == 'revenue_growth':
                # Percentage value
                extracted[key] = float(match) / 100.0
            elif key == 'headcount':
                extracted[key] = int(match)
            elif key == 'capex':
                # CapEx may not have unit suffix
                extracted[key] = _parse_number_text(match)
            elif len(match) == 2:
                # (value, unit) tuple
                extracted[key] = normalize_rmb(match[0], match[1])
            else:
                extracted[key] = _parse_number_text(match)

    # Additional patterns: year-specific data in table-like markdown
    # e.g., | 2023 | 1,234 | 456 |
    table_pattern = r'\|\s*(20[12]\d)\s*\|([^|]+\|)+'
    table_matches = re.finditer(table_pattern, text)
    for m in table_matches:
        year = int(m.group(1))
        if 'year_data' not in extracted:
            extracted['year_data'] = {}
        # Store raw row for later processing
        extracted['year_data'][year] = m.group(0)

    return extracted


# =============================================================================
# Merging
# =============================================================================

def merge_extractions(extractions):
    """Combine data from multiple sources into a single data_dict.

    Priority: xlsx > docx > md (higher confidence from structured sources).

    Args:
        extractions: list of (source_type, extraction_dict) tuples

    Returns:
        Merged data_dict compatible with model_populator
    """
    merged = {
        'company_name': '',
        'industry': '',
        'date': '',
        'analyst': '',
        'currency': 'RMB',
        'segments': [],
        'historical': {},
        'assumptions': {
            'revenue_growth': [],
            'gross_margin_target': [],
            'sga_pct': 0,
            'capex_pct': 0,
            'tax_rate': 0.25,
            'ar_days': 0,
            'ap_days': 0,
            'inventory_days': 0,
            'da_rate': 0,
            'dividend_payout': 0,
            'interest_rate': 0,
        },
        'confidence': {},
    }

    # Priority ordering
    priority = {'xlsx': 3, 'docx': 2, 'md': 1}

    for source_type, extraction in sorted(extractions, key=lambda x: priority.get(x[0], 0)):
        if not extraction:
            continue

        hist = extraction.get('historical', {})
        for year, year_data in hist.items():
            if year not in merged['historical']:
                merged['historical'][year] = {}
            for key, value in year_data.items():
                # Higher priority source overwrites
                merged['historical'][year][key] = value
                # Track confidence
                conf_key = f"{year}.{key}"
                merged['confidence'][conf_key] = {
                    'value': value,
                    'source': extraction.get('source', source_type),
                    'priority': priority.get(source_type, 0),
                }

        # Extract text-based data (from md)
        text_data = extraction.get('text_data', {})
        if text_data:
            if 'gross_margin' in text_data:
                merged['assumptions']['gross_margin_target'] = [text_data['gross_margin']] * 5
            if 'ar_days' in text_data:
                merged['assumptions']['ar_days'] = text_data['ar_days']
            if 'inventory_days' in text_data:
                merged['assumptions']['inventory_days'] = text_data['inventory_days']
            if 'ap_days' in text_data:
                merged['assumptions']['ap_days'] = text_data['ap_days']

    # Derive assumptions from historical data
    _derive_assumptions(merged)

    return merged


def _derive_assumptions(merged):
    """Derive forecast assumptions from historical trends."""
    hist = merged['historical']
    years = sorted(hist.keys())

    if len(years) < 2:
        return

    assumptions = merged['assumptions']

    # Revenue growth (average of last 2-3 years)
    growths = []
    for i in range(1, len(years)):
        prev_rev = hist[years[i - 1]].get('revenue', 0)
        curr_rev = hist[years[i]].get('revenue', 0)
        if prev_rev > 0:
            growths.append((curr_rev - prev_rev) / prev_rev)

    if growths and not assumptions.get('revenue_growth'):
        avg_growth = sum(growths) / len(growths)
        # Taper growth slightly for conservatism
        assumptions['revenue_growth'] = [
            avg_growth * (1 - 0.05 * i) for i in range(5)
        ]

    # Gross margin (average of historical)
    margins = []
    for year in years:
        rev = hist[year].get('revenue', 0)
        cogs = hist[year].get('cogs', 0)
        if rev > 0:
            margins.append(1 - cogs / rev)

    if margins and not assumptions.get('gross_margin_target'):
        avg_margin = sum(margins) / len(margins)
        assumptions['gross_margin_target'] = [avg_margin] * 5

    # SG&A %
    sga_pcts = []
    for year in years:
        rev = hist[year].get('revenue', 0)
        sga = hist[year].get('sga', 0)
        if rev > 0:
            sga_pcts.append(sga / rev)

    if sga_pcts:
        assumptions['sga_pct'] = sum(sga_pcts) / len(sga_pcts)

    # CapEx %
    capex_pcts = []
    for year in years:
        rev = hist[year].get('revenue', 0)
        capex = hist[year].get('capex', 0)
        if rev > 0:
            capex_pcts.append(capex / rev)

    if capex_pcts:
        assumptions['capex_pct'] = sum(capex_pcts) / len(capex_pcts)

    # Tax rate
    tax_rates = []
    for year in years:
        rev = hist[year].get('revenue', 0)
        cogs = hist[year].get('cogs', 0)
        sga = hist[year].get('sga', 0)
        da = hist[year].get('da', 0)
        interest = hist[year].get('interest', 0)
        tax = hist[year].get('tax', 0)
        ebt = rev - cogs - sga - da - interest
        if ebt > 0:
            tax_rates.append(tax / ebt)

    if tax_rates:
        assumptions['tax_rate'] = sum(tax_rates) / len(tax_rates)

    # Working capital days (from latest year)
    latest = years[-1]
    rev = hist[latest].get('revenue', 0)
    cogs = hist[latest].get('cogs', 0)
    ar = hist[latest].get('ar', 0)
    ap = hist[latest].get('ap', 0)
    inventory = hist[latest].get('inventory', 0)

    if rev > 0 and not assumptions.get('ar_days'):
        assumptions['ar_days'] = ar / rev * 365
    if cogs > 0:
        if not assumptions.get('inventory_days'):
            assumptions['inventory_days'] = inventory / cogs * 365
        if not assumptions.get('ap_days'):
            assumptions['ap_days'] = ap / cogs * 365

    # D&A rate
    ppe = hist[latest].get('ppe', 0)
    da = hist[latest].get('da', 0)
    if ppe > 0:
        assumptions['da_rate'] = da / ppe


# =============================================================================
# Utility Functions
# =============================================================================

def normalize_rmb(value, unit):
    """Convert a value with Chinese unit to RMB millions.

    Args:
        value: numeric string (may contain commas)
        unit: Chinese unit string (万, 百万, 千万, 亿)

    Returns:
        float in RMB millions
    """
    multipliers = {'万': 0.01, '百万': 1, '千万': 10, '亿': 100}
    return float(str(value).replace(',', '')) * multipliers.get(unit, 1)


def _to_number(val):
    """Convert a cell value to a float, handling various formats."""
    if val is None:
        return 0
    if isinstance(val, (int, float)):
        return float(val)
    if isinstance(val, str):
        return _parse_number_text(val) or 0
    return 0


def _parse_number_text(text):
    """Parse a number from text, handling commas, parentheses (negative), etc.

    Returns:
        float or None if unparseable
    """
    if not text:
        return None
    if isinstance(text, (int, float)):
        return float(text)

    text = str(text).strip()

    # Handle parentheses as negative
    negative = False
    if text.startswith('(') and text.endswith(')'):
        negative = True
        text = text[1:-1]
    elif text.startswith('-'):
        negative = True
        text = text[1:]

    # Remove commas, spaces, currency symbols
    text = re.sub(r'[,\s¥$€]', '', text)

    # Handle percentage
    if text.endswith('%'):
        text = text[:-1]
        try:
            val = float(text) / 100.0
            return -val if negative else val
        except ValueError:
            return None

    try:
        val = float(text)
        return -val if negative else val
    except ValueError:
        return None


def _try_parse_year(val):
    """Try to parse a year from a cell value.

    Returns:
        int year or None
    """
    if val is None:
        return None
    if isinstance(val, (int, float)):
        v = int(val)
        if 2015 <= v <= 2035:
            return v
        return None
    if isinstance(val, str):
        # Handle "2024E", "2024A", "FY2024", etc.
        match = re.search(r'(20[12]\d)', val)
        if match:
            return int(match.group(1))
    return None


def _merge_text_data(result, text_data):
    """Merge text-extracted data into the result dict."""
    if not text_data:
        return
    # Text data without year context goes into a 'latest' bucket
    # The merge step will handle year assignment
    result['text_data'] = text_data
